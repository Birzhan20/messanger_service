from pathlib import Path
from core.config import settings

try:
    from docx import Document  # python-docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def _read_docx(path: Path) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def _write_docx(path: Path, content: str):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.save(path)


def read_prompt() -> str:
    path = Path(settings.SUPPORT_PROMPT_PATH)

    if not path.exists():
        path.write_text("", encoding="utf-8")
        return ""

    if path.suffix.lower() == ".docx":
        return _read_docx(path)

    return path.read_text(encoding="utf-8")


def write_prompt(content: str):
    path = Path(settings.SUPPORT_PROMPT_PATH)

    if path.suffix.lower() == ".docx":
        _write_docx(path, content)
        return

    path.write_text(content, encoding="utf-8")
