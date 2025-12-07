from pathlib import Path
from core.config import settings
from core.redis import get_redis

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _read_docx(path: Path) -> str:
    """Читает текст из .docx файла."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx не установлен")
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def _write_docx(path: Path, content: str):
    """Записывает текст в .docx файл."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx не установлен")
    doc = Document()
    for line in content.splitlines():
        doc.add_paragraph(line)
    doc.save(path)


def read_prompt() -> str:
    """Читает системный промпт из кэша Redis или файла."""
    redis = get_redis()
    cached = redis.get("support_prompt")
    if cached:
        return cached

    path = Path(settings.SUPPORT_PROMPT_PATH)

    if not path.exists():
        path.write_text("", encoding="utf-8")
        redis.set("support_prompt", "")
        return ""

    if path.suffix.lower() == ".docx":
        content = _read_docx(path)
    else:
        content = path.read_text(encoding="utf-8")

    redis.set("support_prompt", content)
    return content


def write_prompt(content: str):
    """Записывает промпт в файл и обновляет кэш Redis."""
    path = Path(settings.SUPPORT_PROMPT_PATH)

    if path.suffix.lower() == ".docx":
        _write_docx(path, content)
    else:
        path.write_text(content, encoding="utf-8")

    redis = get_redis()
    redis.set("support_prompt", content)


def upload_prompt(file_path: Path):
    """Загружает промпт из файла и сохраняет его."""
    if not file_path.exists():
        raise FileNotFoundError(f"{file_path} не найден")

    if file_path.suffix.lower() in (".txt", ".md"):
        content = file_path.read_text(encoding="utf-8")
    elif file_path.suffix.lower() == ".docx":
        content = _read_docx(file_path)
    else:
        raise ValueError("Поддерживаются только .txt, .md, .docx")

    write_prompt(content)
    return content

